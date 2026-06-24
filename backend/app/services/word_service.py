"""Word document generation for geotechnical reports."""

import logging
import re
import shutil
import unicodedata
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document

from app.core.config import settings
from app.services.resource_manager import resource_manager
from app.utils.common import parse_fecha as _common_parse_fecha, split_project_location as _common_split_location

logger = logging.getLogger(__name__)

INFORME_TEMPLATE_FILENAME = "INFORME BASE SAN PEDRO.docx"
# Municipality that the base template is written for — used to substitute text
# when the base template is applied to a different location.
INFORME_BASE_MUNICIPIO = "SAN PEDRO"

# ── Municipality list for template matching ──────────────────────────────────
_LUGARES_RAW: List[str] = [
    # Departamentos
    "AMAZONAS","ANTIOQUIA","ARAUCA","ATLANTICO","BOLIVAR","BOYACA","CALDAS",
    "CAQUETA","CASANARE","CAUCA","CESAR","CHOCO","CORDOBA","CUNDINAMARCA",
    "GUAINIA","GUAVIARE","HUILA","LA GUAJIRA","MAGDALENA","META","NARINO",
    "NORTE DE SANTANDER","PUTUMAYO","QUINDIO","RISARALDA","SAN ANDRES",
    "SANTANDER","SUCRE","TOLIMA","VALLE DEL CAUCA","VAUPES","VICHADA","BOGOTA",
    # Municipios de Antioquia
    "ABEJORRAL","ABRIAQUI","ALEJANDRIA","AMAGA","ANDES","ANGOSTURA","ANZA",
    "APARTADO","ARBOLETES","ARGELIA","BARBOSA","BELLO","BETANIA","BETULIA",
    "BRICENO","BURITICA","CACERES","CAICEDO","CAMPAMENTO","CANASGORDAS",
    "CARACOLI","CARAMANTA","CAREPA","EL CARMEN","EL CARMEN DE VIBORAL",
    "CAROLINA","CAUCASIA","CHIGORODO","CISNEROS","COCORNA","CONCEPCION",
    "CONCORDIA","COPACABANA","DABEIBA","DON MATIAS","EBEJICO","EL BAGRE",
    "ENTRERRIOS","ENVIGADO","FREDONIA","FRONTINO","GIRALDO","GIRARDOTA",
    "GOMEZ PLATA","GRANADA","GUADALUPE","GUARNE","GUATAPE","HELICONIA",
    "HISPANIA","ITAGUI","ITUANGO","JARDIN","JERICO","LA CEJA","LA ESTRELLA",
    "LA PINTADA","LA UNION","LIBORINA","MACEO","MARINILLA","MEDELLIN",
    "MONTEBELLO","MURINDO","MUTATA","NECHI","NECOCLI","OLAYA","PENOL","PEQUE",
    "PUEBLORRICO","PUERTO BERRIO","PUERTO NARE","PUERTO TRIUNFO","REMEDIOS",
    "EL RETIRO","RIONEGRO","SABANETA","SALGAR","SAN ANDRES DE CUERQUIA",
    "SAN CARLOS","SAN FRANCISCO","SAN JERONIMO","SAN JOSE DE LA MONTANA",
    "SAN JUAN DE URABA","SAN LUIS","SAN PEDRO","SAN RAFAEL","SAN ROQUE",
    "SAN VICENTE","SANTA BARBARA","SANTA FE DE ANTIOQUIA","SANTA ROSA",
    "SANTO DOMINGO","EL SANTUARIO","SANTUARIO","SEGOVIA","SONSON","SOPETRAN",
    "TAMESIS","TARAZA","TARSO","TITIRIBÍ","TOLEDO","TURBO","URAMITA","URRAO",
    "VALDIVIA","VALPARAISO","VEGACHI","VENECIA","VIGIA DEL FUERTE","YALI",
    "YARUMAL","YOLOMBO","YONDO","ZARAGOZA","ANGELOPOLIS",
    # Corregimientos / sectores de Medellín
    "SAN ANTONIO DE PRADO","SANTA ELENA","SAN CRISTOBAL","ALTAVISTA",
    "SAN SEBASTIAN DE PALMITAS","LLANOGRANDE",
    # Barrios frecuentes en informes de Medellín
    "MANRIQUE","LAURELES","BELEN","CASTILLA","BUENOS AIRES","LA AMERICA",
    "VILLA HERMOSA","NIQUIA","EL POBLADO","ARANJUEZ","SAN JAVIER","ROBLEDO",
    # Otros municipios colombianos presentes en los archivos
    "BUGA","PUERTO BOYACA","CURRULAO","PORTACHUELO","PUERTO ESCONDIDO",
    "SAN MARTIN","ALEJANDRIA",
]


def _norm(texto: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFD", texto.upper())
        if unicodedata.category(c) != "Mn"
    )


_LUGARES_NORM: List[str] = [_norm(l) for l in _LUGARES_RAW]

# Module-level cache for the municipios list (populated once per server run)
_MUNICIPIOS_CACHE: Optional[List[Dict]] = None


class WordService:
    """Generate Word copies from the informe template."""

    def __init__(self):
        self.templates_dir = settings.WORD_TEMPLATES_DIR
        self.generated_dir = settings.GENERATED_DIR

    def get_informe_template_path(self) -> Path:
        return self.templates_dir / INFORME_TEMPLATE_FILENAME

    # ── Municipality discovery ────────────────────────────────────────────────

    def get_available_municipios(self) -> List[Dict]:
        """Return unique sorted list of {municipio, filename} for the frontend selector.

        Scans the word templates directory, finds which municipality appears in each
        filename, and returns one best-match file per municipality (fewest extra words
        beyond INFORME + date + place name → cleanest file wins).
        Result is cached for the lifetime of the server process.
        """
        global _MUNICIPIOS_CACHE
        if _MUNICIPIOS_CACHE is not None:
            return _MUNICIPIOS_CACHE

        municipio_files: Dict[int, List[Tuple[Path, int]]] = defaultdict(list)

        # Pre-compile patterns once
        compiled = [
            (idx, re.compile(r"(?:^|[^A-Z])" + re.escape(lugar) + r"(?:[^A-Z]|$)"))
            for idx, lugar in enumerate(_LUGARES_NORM)
        ]

        for path in self.templates_dir.iterdir():
            if path.suffix.lower() not in {".doc", ".docx"}:
                continue
            stem_norm = _norm(path.stem)
            if "INFORME" not in stem_norm:
                continue

            for idx, pattern in compiled:
                if pattern.search(stem_norm):
                    # Score = number of extra words after stripping INFORME, date, place
                    clean = stem_norm
                    clean = clean.replace("INFORME", "")
                    clean = re.sub(r"\d{4}[-/]\d{2}[-/]\d{2}", "", clean)
                    clean = clean.replace(_LUGARES_NORM[idx], "")
                    clean = re.sub(r"[\s\-_.]+", " ", clean).strip()
                    extra = len(clean.split()) if clean else 0
                    municipio_files[idx].append((path, extra))

        seen_norm: set = set()
        result: List[Dict] = []
        for idx in sorted(municipio_files.keys(), key=lambda i: _LUGARES_NORM[i]):
            lugar_norm = _LUGARES_NORM[idx]
            if lugar_norm in seen_norm:
                continue
            seen_norm.add(lugar_norm)
            files = sorted(municipio_files[idx], key=lambda x: x[1])
            best = files[0][0]
            raw = _LUGARES_RAW[idx]
            result.append({
                "municipio": raw.title(),
                "filename": best.name,
            })

        _MUNICIPIOS_CACHE = sorted(result, key=lambda x: _norm(x["municipio"]))
        logger.info("Municipios disponibles en plantillas Word: %d", len(_MUNICIPIOS_CACHE))
        return _MUNICIPIOS_CACHE

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _split_project_location(self, value: Optional[str]) -> Tuple[Optional[str], Optional[str]]:
        """Delega a common.split_project_location — fuente única de verdad."""
        return _common_split_location(value)

    def _parse_fecha(self, fecha_registro) -> Optional[date]:
        """Delega a common.parse_fecha — fuente única de verdad."""
        return _common_parse_fecha(fecha_registro)

    def _format_fecha_for_name(self, fecha_obj: Optional[date], fecha_fallback) -> str:
        if fecha_obj is not None:
            return fecha_obj.isoformat()
        if fecha_fallback is None:
            return ""
        return str(fecha_fallback).strip()

    def build_informe_title(
        self,
        proyecto_ubicacion: str,
        fecha_registro,
    ) -> Tuple[str, str]:
        """Return display title and safe filename (without extension)."""
        _, ubicacion = self._split_project_location(proyecto_ubicacion)
        ubicacion_text = (ubicacion or proyecto_ubicacion or "").strip().upper()
        fecha_obj = self._parse_fecha(fecha_registro)
        fecha_text = self._format_fecha_for_name(fecha_obj, fecha_registro)

        title = " ".join(part for part in ("INFORME", ubicacion_text, fecha_text) if part)
        return title, title

    def _sanitize_docx_filename(self, filename_base: str) -> str:
        text = unicodedata.normalize("NFD", filename_base or "").encode("ascii", "ignore").decode("ascii")
        text = re.sub(r'[<>:"/\\|?*]+', "", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text or "INFORME"

    def _replace_header_text(self, document: Document, title: str) -> None:
        for section in document.sections:
            header = section.header
            if not header.paragraphs:
                header.add_paragraph(title)
                continue

            first_paragraph = header.paragraphs[0]
            if first_paragraph.runs:
                first_paragraph.runs[0].text = title
                for run in first_paragraph.runs[1:]:
                    run.text = ""
            else:
                first_paragraph.text = title

            for extra_paragraph in header.paragraphs[1:]:
                extra_paragraph.text = ""

    # ── Placeholder replacement ───────────────────────────────────────────────

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        """Replace placeholder strings in a paragraph, handling run-split text.

        python-docx can split a single word across multiple runs due to spell-check
        or formatting changes. We join all run texts, replace, then put the result
        back into the first run to preserve its character formatting.
        """
        if not paragraph.runs:
            return
        original = "".join(run.text for run in paragraph.runs)
        if not original:
            return
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, str(new))
        if updated != original:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""

    def _process_body(self, container, replacements: Dict[str, str]) -> None:
        """Recursively process paragraphs and tables in a document body or cell."""
        for paragraph in container.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_body(cell, replacements)

    def _build_replacements(
        self,
        proyecto_ubicacion: str,
        fecha_registro,
        municipio_word: Optional[str] = None,
        cliente: Optional[str] = None,
        sondeo: Optional[str] = None,
        pisos: Optional[int] = None,
        perforaciones: Optional[List[Dict]] = None,
        clasificacion_suelo: Optional[str] = None,
    ) -> Dict[str, str]:
        """Build the placeholder→value map used to fill the Word document.

        Placeholders follow the {{NOMBRE}} convention. Add these tags to the
        Word base template wherever dynamic data should appear.
        """
        fecha_obj = self._parse_fecha(fecha_registro)
        _MESES_ES = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
        ]
        if fecha_obj:
            fecha_larga = f"{fecha_obj.day} de {_MESES_ES[fecha_obj.month - 1]} de {fecha_obj.year}"
            fecha_corta = fecha_obj.strftime("%d/%m/%Y")
            año = str(fecha_obj.year)
        else:
            fecha_larga = str(fecha_registro or "")
            fecha_corta = fecha_larga
            año = ""

        _, ubicacion = self._split_project_location(proyecto_ubicacion)
        ubicacion_text = (ubicacion or proyecto_ubicacion or "").strip().upper()
        municipio_text = (municipio_word or "").strip().upper()

        # Build a plain-text description of soil layers for insertion in the document
        descripcion_capas = ""
        profundidad_total = ""
        if perforaciones:
            partes: List[str] = []
            max_prof = 0.0
            for i, p in enumerate(perforaciones, 1):
                desc = (p.get("descripcion_suelo") or p.get("tipo_suelo_principal") or "").strip()
                prof = p.get("profundidad_z")
                color = (p.get("color_predominante") or "").strip().lower()
                try:
                    v = float(str(prof).replace(",", "."))
                    if v > max_prof:
                        max_prof = v
                except Exception:
                    pass
                if desc:
                    parte = f"Capa {i}: {desc}"
                    if color:
                        parte += f", color {color}"
                    if prof:
                        parte += f", hasta {prof}m de profundidad"
                    partes.append(parte)
            descripcion_capas = ". ".join(partes) if partes else ""
            profundidad_total = f"{max_prof:.1f}" if max_prof > 0 else ""

        return {
            "{{PROYECTO}}": (proyecto_ubicacion or "").upper(),
            "{{UBICACION}}": ubicacion_text,
            "{{CLIENTE}}": (cliente or "").upper(),
            "{{FECHA}}": fecha_larga,
            "{{FECHA_CORTA}}": fecha_corta,
            "{{AÑO}}": año,
            "{{MUNICIPIO}}": municipio_text,
            "{{SONDEO}}": sondeo or "",
            "{{PISOS}}": str(pisos) if pisos is not None else "",
            "{{N_PISOS}}": str(pisos) if pisos is not None else "",
            "{{CLASIFICACION}}": (clasificacion_suelo or "").upper(),
            "{{CAPAS_SUELO}}": descripcion_capas,
            "{{PROFUNDIDAD}}": profundidad_total,
        }

    def _replace_placeholders_in_document(
        self,
        document: Document,
        replacements: Dict[str, str],
        substitute_municipio: Optional[str] = None,
    ) -> None:
        """Apply all placeholder replacements throughout the Word document.

        If *substitute_municipio* is provided, every occurrence of the base template's
        municipality name (SAN PEDRO) in the body text is also replaced so that the
        document reads as if it were written for the target location.
        """
        all_reps = dict(replacements)

        if substitute_municipio:
            sub_upper = substitute_municipio.strip().upper()
            base_upper = INFORME_BASE_MUNICIPIO.upper()
            if _norm(sub_upper) != _norm(base_upper):
                all_reps[base_upper] = sub_upper

        # Document body (paragraphs + tables, recursively)
        self._process_body(document, all_reps)

        # Headers and footers
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, all_reps)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._process_body(cell, all_reps)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, all_reps)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._process_body(cell, all_reps)

    # ── Main generation method ────────────────────────────────────────────────

    def generate_informe(
        self,
        project_id: str,
        proyecto_ubicacion: str,
        fecha_registro,
        municipio_word: Optional[str] = None,
        template_filename: Optional[str] = None,
        cliente: Optional[str] = None,
        sondeo: Optional[str] = None,
        pisos: Optional[int] = None,
        perforaciones: Optional[List[Dict]] = None,
        clasificacion_suelo: Optional[str] = None,
    ) -> Tuple[Path, str]:
        """Copy an informe template, fill in data placeholders and update the header.

        Template selection logic:
        - template_filename set  → use that specific Word file as content source
        - municipio_word set     → use its value in the title / output filename
        - Both set               → use template file for content, municipio_word for title
        - Neither set            → base template + title built from proyecto_ubicacion

        When the base template (INFORME BASE SAN PEDRO) is selected and a different
        municipality is specified, every occurrence of "SAN PEDRO" in the document
        body is replaced with the target municipality name so the informe reads
        correctly without a dedicated template file.
        """
        # ── Determine source template ─────────────────────────────────────────
        if template_filename:
            candidate = self.templates_dir / Path(template_filename).name
            template_path = candidate if candidate.exists() else self.get_informe_template_path()
        else:
            template_path = self.get_informe_template_path()

        # If the selected or default template is missing, fall back to any available .docx
        if not template_path.exists():
            fallback = next(
                (p for p in sorted(self.templates_dir.glob("*.docx")) if p.is_file()),
                None,
            )
            if fallback is None:
                raise FileNotFoundError(f"No se encontró la plantilla Word: {template_path}")
            logger.warning("Plantilla '%s' no encontrada; usando '%s' como base.", template_path.name, fallback.name)
            template_path = fallback

        # ── Build title and archive name ──────────────────────────────────────
        fecha_obj = self._parse_fecha(fecha_registro)
        fecha_text = self._format_fecha_for_name(fecha_obj, fecha_registro)

        if municipio_word and municipio_word.strip():
            mun_text = municipio_word.strip().upper()
            title = " ".join(p for p in ("INFORME", mun_text, fecha_text) if p)
        else:
            title, _ = self.build_informe_title(proyecto_ubicacion, fecha_registro)

        safe_base = self._sanitize_docx_filename(title)
        output_path = self.generated_dir / f"{project_id}_{safe_base}.docx"
        archive_name = f"{safe_base}.docx"

        # ── Copy, fill placeholders, update header, save ──────────────────────
        shutil.copy2(template_path, output_path)
        document = Document(output_path)

        # Build placeholder replacements from all available data
        replacements = self._build_replacements(
            proyecto_ubicacion=proyecto_ubicacion,
            fecha_registro=fecha_registro,
            municipio_word=municipio_word,
            cliente=cliente,
            sondeo=sondeo,
            pisos=pisos,
            perforaciones=perforaciones,
            clasificacion_suelo=clasificacion_suelo,
        )

        # When using the base template for a different municipality, substitute
        # the embedded "SAN PEDRO" references with the actual target municipality.
        using_base = (template_path.name == INFORME_TEMPLATE_FILENAME)
        substitute = None
        if using_base and municipio_word and _norm(municipio_word) != _norm(INFORME_BASE_MUNICIPIO):
            substitute = municipio_word

        self._replace_placeholders_in_document(document, replacements, substitute_municipio=substitute)

        # Header title is set last so it takes precedence over any placeholder replacement
        self._replace_header_text(document, title)

        document.save(output_path)

        # Aplicar recursos visuales registrados para documentos Word.
        # Cuando se declaren imágenes Word en resource_manager, este llamado
        # las insertará automáticamente sin cambiar este código.
        try:
            resource_manager.apply_to_word(output_path, "word_informe")
        except Exception:
            logger.debug(
                "No se pudieron aplicar recursos visuales al informe Word", exc_info=True
            )

        logger.info(
            "Informe Word generado: %s (fuente: %s, sustitución municipio: %s)",
            output_path.name, template_path.name, substitute or "ninguna",
        )
        return output_path, archive_name


word_service = WordService()
